"""
Document readers for RAG ingestion.

One entry point — read_document() — dispatches on file extension.
PDFs use PyMuPDF (5-10x faster than PyPDFLoader) with parallel page
extraction for large files, falling back to sequential PyMuPDF and then
PyPDFLoader if unavailable.
"""
import json
from concurrent.futures import ProcessPoolExecutor, as_completed
from pathlib import Path
from typing import Callable, List, Optional

# Parallel extraction only pays off past this page count; below it the
# process-pool spawn cost outweighs the parsing work.
_PARALLEL_PDF_MIN_PAGES = 100
_PAGES_PER_BATCH = 20


def read_document(path: Path, progress_callback: Optional[Callable[[str, float], None]] = None) -> str:
    """Read a document's text content based on its file type."""
    suffix = path.suffix.lower()

    if suffix == '.pdf':
        return read_pdf(path, progress_callback=progress_callback)

    if suffix == '.json':
        with open(path, 'r', encoding='utf-8') as f:
            return json.dumps(json.load(f), indent=2)

    if suffix == '.csv':
        import pandas as pd
        return pd.read_csv(path).to_string()

    if suffix in ('.xlsx', '.xls'):
        import pandas as pd
        excel_file = pd.ExcelFile(path)
        parts = []
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(path, sheet_name=sheet_name)
            parts.append(f"[Sheet: {sheet_name}]")
            parts.append(df.to_string())
        return '\n\n'.join(parts)

    if suffix == '.docx':
        from docx import Document
        doc = Document(path)
        parts = [para.text for para in doc.paragraphs if para.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(' | '.join(cell.text.strip() for cell in row.cells))
        return '\n'.join(parts)

    # .txt, .md, and anything else readable as UTF-8 text
    with open(path, 'r', encoding='utf-8') as f:
        return f.read()


def read_pdf(path: Path, progress_callback: Optional[Callable[[str, float], None]] = None) -> str:
    """Extract text from a PDF, fastest available path first."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        from langchain_community.document_loaders import PyPDFLoader
        pages = PyPDFLoader(str(path)).load()
        return '\n'.join(page.page_content for page in pages)

    doc = fitz.open(str(path))
    total_pages = len(doc)
    doc.close()

    if total_pages >= _PARALLEL_PDF_MIN_PAGES:
        try:
            return _read_pdf_parallel(path, total_pages, progress_callback)
        except Exception as e:
            print(f"[RAG] Parallel PDF extraction failed ({e}), reading sequentially")

    doc = fitz.open(str(path))
    text = "\n\n".join(page.get_text() for page in doc)
    doc.close()
    return text


def _read_pdf_parallel(
    path: Path,
    total_pages: int,
    progress_callback: Optional[Callable[[str, float], None]] = None,
) -> str:
    """Extract PDF text with one process per page batch (all CPU cores)."""
    batches = [
        (str(path), start, min(start + _PAGES_PER_BATCH, total_pages))
        for start in range(0, total_pages, _PAGES_PER_BATCH)
    ]

    parts: List[tuple] = []
    with ProcessPoolExecutor() as executor:
        futures = {
            executor.submit(_extract_page_batch, *batch): i
            for i, batch in enumerate(batches)
        }
        for done, future in enumerate(as_completed(futures), 1):
            parts.append((futures[future], future.result()))
            if progress_callback:
                progress_callback(f"Extracted batch {done}/{len(batches)}", 100.0 * done / len(batches))

    parts.sort(key=lambda x: x[0])
    return "\n\n".join(text for _, text in parts)


def _extract_page_batch(pdf_path: str, page_start: int, page_end: int) -> str:
    """Worker: extract text from a range of pages. Must stay module-level picklable."""
    import fitz
    doc = fitz.open(pdf_path)
    text = "\n\n".join(doc[n].get_text() for n in range(page_start, min(page_end, len(doc))))
    doc.close()
    return text


def load_chunks_from_disk(chunks_file: Path) -> List[str]:
    """Load a document's chunk list from its spill file (legacy large uploads)."""
    with open(chunks_file, 'r', encoding='utf-8') as f:
        return json.load(f)
