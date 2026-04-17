"""
RAG (Retrieval-Augmented Generation) Tool
Uses FAISS for vector similarity search with configurable embeddings
"""
import json
import os
import time
import hashlib
from pathlib import Path
from typing import Dict, List, Optional, Any
from datetime import datetime
import numpy as np

import config
from backend.utils.prompts_log_append import log_to_prompts_file

# ---------------------------------------------------------------------------
# Process-level singletons — loaded once, reused across all requests
# ---------------------------------------------------------------------------
# _GLOBAL_EMBEDDING_MODEL is the SINGLE source of truth for the embedding
# model in this worker process. Both BaseRAGTool (this file) and
# EnhancedRAGTool (enhanced_tool.py) read/write it through
# get_global_embedding_model(). Never instantiate SentenceTransformer
# anywhere else — duplicates would stack another ~2.3 GB of VRAM per copy.
_GLOBAL_EMBEDDING_MODEL = None   # SentenceTransformer (GPU-resident)
_GLOBAL_FAISS_CACHE: Dict[str, Any] = {}  # "path:mtime" → faiss.Index


def get_global_embedding_model():
    """Load-or-return the shared SentenceTransformer singleton for this process.

    Call from anywhere that needs the embedding model. First call loads weights
    onto `config.RAG_EMBEDDING_DEVICE`; every subsequent call returns the exact
    same object. This guarantees one and only one embedding-model VRAM stack
    per worker process, regardless of how many RAG tool instances are created.
    """
    global _GLOBAL_EMBEDDING_MODEL
    if _GLOBAL_EMBEDDING_MODEL is not None:
        return _GLOBAL_EMBEDDING_MODEL

    try:
        from sentence_transformers import SentenceTransformer
    except ImportError:
        raise ImportError(
            "sentence-transformers not installed. "
            "Install with: pip install sentence-transformers"
        )

    # Load in FP16 to halve model-weight VRAM (~2.3 GB → ~1.15 GB for bge-m3).
    # Accuracy impact on bge-m3 is near-zero in practice.
    model_kwargs = {}
    if config.RAG_EMBEDDING_DEVICE == "cuda":
        try:
            import torch
            model_kwargs["torch_dtype"] = torch.float16
            print(f"[RAG] Loading embedding model in FP16 (half VRAM footprint)")
        except ImportError:
            pass

    print(f"[RAG] Loading embedding model: {config.RAG_EMBEDDING_MODEL}")
    _GLOBAL_EMBEDDING_MODEL = SentenceTransformer(
        config.RAG_EMBEDDING_MODEL,
        device=config.RAG_EMBEDDING_DEVICE,
        model_kwargs=model_kwargs,
    )
    dim = _GLOBAL_EMBEDDING_MODEL.get_sentence_embedding_dimension()
    print(f"[RAG] Embedding model loaded - dimension: {dim}")

    # Diagnostic: show VRAM state after model load so the operator can confirm
    # FP16 savings and verify the allocator config is in effect.
    try:
        import torch
        if torch.cuda.is_available():
            alloc = torch.cuda.memory_allocated() / 1e9
            resv = torch.cuda.memory_reserved() / 1e9
            alloc_conf = os.environ.get("PYTORCH_CUDA_ALLOC_CONF", "(not set)")
            print(f"[RAG] VRAM after model load: alloc={alloc:.2f} GB  "
                  f"reserved={resv:.2f} GB  "
                  f"PYTORCH_CUDA_ALLOC_CONF={alloc_conf}")
    except ImportError:
        pass

    return _GLOBAL_EMBEDDING_MODEL


def _release_cuda_cache(context: str = ""):
    """Return freed CUDA blocks to the driver so nvidia-smi drops back down.

    Cost is one cudaStreamSynchronize + freelist bookkeeping (~10-30 ms).
    Irrelevant in upload paths (seconds/minutes), acceptable as a conditional
    call on the retrieve hot-path when fragmentation is severe.
    """
    try:
        import torch
        if torch.cuda.is_available():
            before = torch.cuda.memory_reserved() / (1024 ** 3)
            torch.cuda.empty_cache()
            after = torch.cuda.memory_reserved() / (1024 ** 3)
            freed = before - after
            if freed > 0.01:
                print(f"[RAG] CUDA cache released ({context}): "
                      f"reserved {before:.2f} → {after:.2f} GiB (freed {freed:.2f} GiB)")
    except ImportError:
        pass


def _maybe_release_cuda_cache(threshold_gib: float = 4.0):
    """Release CUDA cache only when fragmentation exceeds *threshold_gib*.

    Fragmentation = memory_reserved - memory_allocated.  Calling empty_cache()
    on every retrieve() would add 10-30 ms of user-visible latency for no gain
    when fragmentation is small.  This conditional variant keeps the hot path
    fast while still bounding runaway reserved memory between uploads.
    """
    try:
        import torch
        if torch.cuda.is_available():
            reserved = torch.cuda.memory_reserved()
            allocated = torch.cuda.memory_allocated()
            frag_gib = (reserved - allocated) / (1024 ** 3)
            if frag_gib > threshold_gib:
                _release_cuda_cache(
                    context=f"retrieve fragmentation {frag_gib:.1f} GiB > {threshold_gib} GiB threshold"
                )
    except ImportError:
        pass


def _ensure_chunk_lookup(metadata: Dict[str, Any]) -> Dict[str, Dict[str, Any]]:
    lookup = metadata.get("chunk_lookup")
    if isinstance(lookup, dict):
        return lookup
    metadata["chunk_lookup"] = {}
    return metadata["chunk_lookup"]


def _set_chunk_lookup_for_doc(
    metadata: Dict[str, Any],
    doc_id: str,
    chunk_indices: List[int],
):
    lookup = _ensure_chunk_lookup(metadata)
    for local_idx, global_idx in enumerate(chunk_indices):
        lookup[str(global_idx)] = {
            "doc_id": doc_id,
            "chunk_index": local_idx,
        }


def _rebuild_chunk_lookup(metadata: Dict[str, Any]) -> Dict[str, Dict[str, Any]]:
    lookup: Dict[str, Dict[str, Any]] = {}
    for doc_id, doc_meta in metadata.get("documents", {}).items():
        for local_idx, global_idx in enumerate(doc_meta.get("chunk_indices", [])):
            lookup[str(global_idx)] = {
                "doc_id": doc_id,
                "chunk_index": local_idx,
            }
    metadata["chunk_lookup"] = lookup
    return lookup


class RAGTool:
    """
    RAG tool with FAISS vector database
    Per-user collections with document upload and retrieval
    """

    def __init__(self, username: str):
        """
        Initialize RAG tool for a user

        Args:
            username: Username for collection isolation
        """
        self.username = username
        self.user_docs_dir = config.RAG_DOCUMENTS_DIR / username
        self.user_index_dir = config.RAG_INDEX_DIR / username
        self.user_metadata_dir = config.RAG_METADATA_DIR / username

        # Create directories
        self.user_docs_dir.mkdir(parents=True, exist_ok=True)
        self.user_index_dir.mkdir(parents=True, exist_ok=True)
        self.user_metadata_dir.mkdir(parents=True, exist_ok=True)

        # Load embedding model
        self.embedding_model = None  # Lazy load
        self.embedding_dim = None

    def _load_embedding_model(self):
        """Bind the process-level singleton embedding model to this instance."""
        self.embedding_model = get_global_embedding_model()
        self.embedding_dim = self.embedding_model.get_sentence_embedding_dimension()

    def create_collection(self, collection_name: str) -> Dict[str, Any]:
        """
        Create a new document collection

        Args:
            collection_name: Name of the collection

        Returns:
            Result dictionary
        """
        collection_dir = self.user_docs_dir / collection_name
        index_path = self.user_index_dir / f"{collection_name}.index"
        metadata_path = self.user_metadata_dir / f"{collection_name}.json"

        # Check if collection already exists with valid metadata
        if metadata_path.exists():
            try:
                with open(metadata_path, 'r', encoding='utf-8') as f:
                    existing_metadata = json.load(f)
                return {
                    "success": False,
                    "error": f"Collection '{collection_name}' already exists with {len(existing_metadata.get('documents', {}))} documents"
                }
            except Exception:
                # Metadata file exists but is corrupted, recreate it
                print(f"[RAG] Warning: Corrupted metadata found for '{collection_name}', recreating...")

        # Create collection directory
        collection_dir.mkdir(parents=True, exist_ok=True)

        # Create empty metadata
        metadata = {
            "collection_name": collection_name,
            "created_at": time.time(),
            "documents": {},
            "chunk_count": 0,
            "chunk_lookup": {},
        }

        try:
            with open(metadata_path, 'w', encoding='utf-8') as f:
                json.dump(metadata, f, indent=2)
        except Exception as e:
            return {
                "success": False,
                "error": f"Failed to create metadata file: {str(e)}"
            }

        return {
            "success": True,
            "collection_name": collection_name,
            "path": str(collection_dir)
        }

    def upload_document(
        self,
        collection_name: str,
        document_path: str,
        document_content: Optional[str] = None,
        document_name: Optional[str] = None,
        use_optimized: bool = True,
        progress_callback: Optional[Any] = None
    ) -> Dict[str, Any]:
        """
        Upload and index a document

        Args:
            collection_name: Collection to add to
            document_path: Path to document or document name
            document_content: Document content (if providing directly)
            document_name: Optional override for document name (useful when uploading from temp files)
            use_optimized: Use optimized uploader for PDFs (5-10x faster)
            progress_callback: Optional callback(message, progress_pct) for progress updates

        Returns:
            Upload result
        """
        # Check if this is a PDF and we should use optimized path
        doc_path = Path(document_path)
        if (use_optimized and 
            document_content is None and 
            doc_path.exists() and 
            doc_path.suffix.lower() == '.pdf'):
            
            print(f"[RAG] Using optimized PDF uploader for: {doc_path.name}")
            return self._upload_pdf_optimized(
                collection_name=collection_name,
                pdf_path=doc_path,
                document_name=document_name,
                progress_callback=progress_callback
            )
        
        # Standard upload path for non-PDFs or when optimization disabled
        vram_before = self._vram_mb()
        self._reset_peak_vram()
        self._load_embedding_model()

        collection_dir = self.user_docs_dir / collection_name
        metadata_path = self.user_metadata_dir / f"{collection_name}.json"

        # Create collection if it doesn't exist
        if not collection_dir.exists():
            collection_dir.mkdir(parents=True, exist_ok=True)
            print(f"[RAG] Created new collection directory: {collection_dir}")

        # Load or create metadata
        if metadata_path.exists():
            try:
                with open(metadata_path, 'r', encoding='utf-8') as f:
                    metadata = json.load(f)
                print(f"[RAG] Loaded existing metadata ({len(metadata.get('documents', {}))} documents)")
            except Exception as e:
                print(f"[RAG] Warning: Failed to load metadata, creating new: {e}")
                metadata = {
                    "collection_name": collection_name,
                    "created_at": time.time(),
                    "documents": {},
                    "chunk_count": 0,
                    "chunk_lookup": {},
                }
        else:
            print(f"[RAG] Creating new metadata file")
            metadata = {
                "collection_name": collection_name,
                "created_at": time.time(),
                "documents": {},
                "chunk_count": 0,
                "chunk_lookup": {},
            }

        _ensure_chunk_lookup(metadata)

        # Read document
        if document_content is None:
            doc_path = Path(document_path)
            if not doc_path.exists():
                return {"success": False, "error": "Document file not found"}

            if doc_path.suffix not in config.RAG_SUPPORTED_FORMATS:
                return {
                    "success": False,
                    "error": f"Unsupported format: {doc_path.suffix}"
                }

            document_content = self._read_document(doc_path)
            doc_name = document_name if document_name else doc_path.name
        else:
            doc_name = document_name if document_name else Path(document_path).name

        print(f"[RAG] Processing document: {doc_name}")

        # Chunk document
        print(f"[RAG] Chunking document...")
        chunks = self._chunk_text(document_content)
        print(f"[RAG] Created {len(chunks)} chunks")

        # Generate embeddings (normalized for cosine similarity with IndexFlatIP)
        print(f"[RAG] Generating embeddings...")
        embeddings = self.embedding_model.encode(
            chunks,
            batch_size=config.RAG_EMBEDDING_BATCH_SIZE,
            show_progress_bar=False,
            normalize_embeddings=True
        )
        print(f"[RAG] Generated {len(embeddings)} embeddings")

        # Load or create index
        print(f"[RAG] Loading/creating FAISS index...")
        index = self._load_or_create_index(collection_name, self.embedding_dim)
        print(f"[RAG] Index loaded (current size: {index.ntotal} vectors)")

        # Add to index
        start_idx = index.ntotal
        print(f"[RAG] Adding embeddings to index (indices {start_idx} to {start_idx + len(embeddings) - 1})...")
        index.add(np.array(embeddings).astype('float32'))
        print(f"[RAG] Index now has {index.ntotal} vectors")
        vram_peak = self._peak_vram_mb()
        vram_after = self._vram_mb()

        # Update metadata (include timestamp in hash to avoid collision on re-upload)
        upload_time = time.time()
        doc_id = hashlib.md5(f"{doc_name}:{upload_time}".encode()).hexdigest()
        chunk_indices = list(range(start_idx, index.ntotal))
        metadata["documents"][doc_id] = {
            "name": doc_name,
            "path": str(document_path),
            "chunk_indices": chunk_indices,
            "chunks": chunks,
            "uploaded_at": upload_time
        }
        metadata["chunk_count"] = index.ntotal
        _set_chunk_lookup_for_doc(metadata, doc_id, chunk_indices)

        # Save index and metadata
        print(f"[RAG] Saving index to disk...")
        try:
            self._save_index(index, collection_name)
            print(f"[RAG] Index saved successfully")
        except Exception as e:
            return {
                "success": False,
                "error": f"Failed to save index: {str(e)}"
            }

        print(f"[RAG] Saving metadata to disk...")
        try:
            with open(metadata_path, 'w', encoding='utf-8') as f:
                json.dump(metadata, f, indent=2)
            print(f"[RAG] Metadata saved successfully")
        except Exception as e:
            return {
                "success": False,
                "error": f"Failed to save metadata: {str(e)}"
            }

        print(f"[RAG] Upload complete!")
        _release_cuda_cache("BaseRAGTool.upload_document")
        result = {
            "success": True,
            "document_name": doc_name,
            "chunks_created": len(chunks),
            "total_chunks": index.ntotal
        }
        if vram_before is not None:
            result["vram_before_mb"] = round(vram_before, 1)
            result["vram_after_mb"] = round(vram_after, 1) if vram_after is not None else None
            result["vram_peak_mb"] = round(vram_peak, 1) if vram_peak is not None else None
            result["vram_delta_mb"] = round((vram_after - vram_before), 1) if vram_after is not None else None
        return result

    def retrieve(
        self,
        collection_name: str,
        query: str,
        max_results: Optional[int] = None
    ) -> Dict[str, Any]:
        """
        Retrieve relevant documents for a query

        Args:
            collection_name: Collection to search
            query: Search query
            max_results: Maximum results to return

        Returns:
            Retrieved documents
        """
        # Batch log writes to reduce FileLock acquisitions
        _log_lines = [
            "\n\n", "=" * 80, "TOOL EXECUTION: rag", "=" * 80,
            f"Timestamp: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", "",
            "INPUT:", f"  Username: {self.username}",
            f"  Collection: {collection_name}", f"  Query: {query}",
            f"  Max Results: {max_results or 'default'}",
        ]

        print("\n" + "=" * 80)
        print("[RAG TOOL] retrieve() called")
        print("=" * 80)
        print(f"Username: {self.username}")
        print(f"Collection: {collection_name}")
        print(f"Query: {query}")
        print(f"Max results: {max_results or 'default'}")

        print(f"\n[RAG] Loading embedding model...")
        start_time = time.time()
        self._load_embedding_model()
        load_time = time.time() - start_time
        print(f"[RAG] [OK] Embedding model loaded ({load_time:.2f}s)")
        print(f"  Model: {config.RAG_EMBEDDING_MODEL}")
        print(f"  Dimension: {self.embedding_dim}")

        metadata_path = self.user_metadata_dir / f"{collection_name}.json"
        print(f"\n[RAG] Checking collection...")
        print(f"  Metadata path: {metadata_path}")

        if not metadata_path.exists():
            print(f"[RAG] [ERROR] Collection not found!")

            _log_lines.extend(["", "OUTPUT:", "  Status: ERROR",
                f"  Error: Collection '{collection_name}' does not exist", "", "=" * 80])
            log_to_prompts_file("\n".join(_log_lines))

            return {
                "success": False,
                "error": f"Collection '{collection_name}' does not exist"
            }
        
        print(f"[RAG] [OK] Collection found")

        # Load index and metadata
        print(f"\n[RAG] Loading FAISS index...")
        index = self._load_index(collection_name)
        if index is None:
            print(f"[RAG] [ERROR] Index not found!")

            _log_lines.extend(["", "OUTPUT:", "  Status: ERROR",
                f"  Error: No index found for collection '{collection_name}'", "", "=" * 80])
            log_to_prompts_file("\n".join(_log_lines))

            return {
                "success": False,
                "error": f"No index found for collection '{collection_name}'"
            }
        print(f"[RAG] [OK] Index loaded")
        print(f"  Total vectors: {index.ntotal}")

        print(f"\n[RAG] Loading metadata...")
        with open(metadata_path, 'r', encoding='utf-8') as f:
            metadata = json.load(f)
        print(f"[RAG] [OK] Metadata loaded")
        print(f"  Documents: {len(metadata.get('documents', {}))}")
        print(f"  Chunks: {metadata.get('chunk_count', 0)}")

        chunk_lookup = metadata.get("chunk_lookup")
        if not isinstance(chunk_lookup, dict) or len(chunk_lookup) != metadata.get("chunk_count", 0):
            chunk_lookup = _rebuild_chunk_lookup(metadata)
            try:
                with open(metadata_path, 'w', encoding='utf-8') as f:
                    json.dump(metadata, f, indent=2)
            except Exception:
                pass

        if index.ntotal == 0:
            print(f"[RAG] [WARNING] Collection is empty")
            return {
                "success": True,
                "documents": [],
                "message": "Collection is empty"
            }

        # Generate query embedding
        print(f"\n[RAG] Generating query embedding...")
        embed_start = time.time()
        query_embedding = self.embedding_model.encode(
            [config.RAG_QUERY_PREFIX + query],
            normalize_embeddings=True
        )[0]
        embed_time = time.time() - embed_start
        print(f"[RAG] [OK] Query embedding generated ({embed_time:.2f}s)")
        print(f"  Embedding shape: {query_embedding.shape}")

        # Search
        k = min(max_results or config.RAG_MAX_RESULTS, index.ntotal)
        print(f"\n[RAG] Searching FAISS index...")
        print(f"  K (results to retrieve): {k}")
        
        search_start = time.time()
        distances, indices = index.search(
            np.array([query_embedding]).astype('float32'),
            k
        )
        search_time = time.time() - search_start
        print(f"[RAG] [OK] Search completed ({search_time:.2f}s)")
        print(f"  Distances: {distances[0].tolist()}")
        print(f"  Indices: {indices[0].tolist()}")

        # Retrieve chunks with context window
        print(f"\n[RAG] Retrieving document chunks (context_window={config.RAG_CONTEXT_WINDOW})...")

        results = []
        for i, (dist, idx) in enumerate(zip(distances[0], indices[0])):
            # Skip invalid FAISS indices (FAISS returns -1 for empty slots when k > ntotal)
            if idx < 0:
                continue

            ref = chunk_lookup.get(str(int(idx)))
            if not ref:
                continue

            doc_meta = metadata["documents"].get(ref["doc_id"])
            if not doc_meta:
                continue

            chunk_local_idx = int(ref["chunk_index"])
            if config.RAG_SIMILARITY_METRIC == "cosine":
                score = float(dist)
            else:
                score = float(1 / (1 + dist))

            doc_chunks = self._get_document_chunks(doc_meta)
            if not doc_chunks or chunk_local_idx >= len(doc_chunks):
                continue
            chunk_text = doc_chunks[chunk_local_idx]

            # Build context window from neighboring chunks
            total_chunks = len(doc_chunks)
            window = config.RAG_CONTEXT_WINDOW
            start_ctx = max(0, chunk_local_idx - window)
            end_ctx = min(total_chunks, chunk_local_idx + window + 1)
            chunk_with_context = "\n---\n".join(doc_chunks[start_ctx:end_ctx])

            results.append({
                "document": doc_meta["name"],
                "chunk": chunk_with_context,
                "score": score,
                "chunk_index": chunk_local_idx
            })

            print(f"  Result {i+1}: {doc_meta['name']} chunk {chunk_local_idx} (score: {score:.3f})")
            print(f"    Context: chunks {start_ctx}-{end_ctx-1}, Preview: {chunk_text[:100]}...")

        # Filter out low-relevance results
        pre_filter_count = len(results)
        results = [r for r in results if r["score"] >= config.RAG_MIN_SCORE_THRESHOLD]
        filtered_count = pre_filter_count - len(results)
        if filtered_count > 0:
            print(f"\n[RAG] Filtered {filtered_count} results below score threshold ({config.RAG_MIN_SCORE_THRESHOLD})")

        print(f"\n[RAG] [OK] Retrieved {len(results)} results")
        total_time = time.time() - start_time
        print(f"[RAG] Total time: {total_time:.2f}s")

        # Flush batched log in a single FileLock acquisition
        _log_lines.extend(["", "OUTPUT:", "  Status: SUCCESS",
            f"  Results Found: {len(results)}", f"  Execution Time: {total_time:.2f}s",
            "", "RESULTS:"])
        for i, result in enumerate(results, 1):
            chunk_preview = result['chunk'][:200] if len(result['chunk']) > 200 else result['chunk']
            _log_lines.extend(["", f"  Result {i}:", f"    Document: {result['document']}",
                f"    Chunk Index: {result['chunk_index']}", f"    Score: {result['score']:.3f}",
                f"    Content: {chunk_preview}..."])
        _log_lines.extend(["", "=" * 80])
        log_to_prompts_file("\n".join(_log_lines))

        _maybe_release_cuda_cache()

        return {
            "success": True,
            "documents": results,
            "query": query,
            "num_results": len(results)
        }

    def list_collections(self) -> Dict[str, Any]:
        """
        List all collections for user

        Returns:
            List of collections
        """
        collections = []

        for metadata_file in self.user_metadata_dir.glob("*.json"):
            try:
                with open(metadata_file, 'r', encoding='utf-8') as f:
                    metadata = json.load(f)

                collections.append({
                    "name": metadata["collection_name"],
                    "documents": len(metadata["documents"]),
                    "chunks": metadata["chunk_count"],
                    "created_at": metadata["created_at"]
                })
            except Exception:
                continue

        return {
            "success": True,
            "collections": collections
        }

    def delete_collection(self, collection_name: str) -> Dict[str, Any]:
        """
        Delete a collection

        Args:
            collection_name: Collection to delete

        Returns:
            Result dictionary
        """
        import shutil

        collection_dir = self.user_docs_dir / collection_name
        index_path = self.user_index_dir / f"{collection_name}.index"
        metadata_path = self.user_metadata_dir / f"{collection_name}.json"

        if not metadata_path.exists():
            return {
                "success": False,
                "error": f"Collection '{collection_name}' does not exist"
            }

        # Delete all files
        if collection_dir.exists():
            shutil.rmtree(collection_dir)
        if index_path.exists():
            index_path.unlink()
        if metadata_path.exists():
            metadata_path.unlink()

        return {
            "success": True,
            "collection_name": collection_name
        }

    def list_documents(self, collection_name: str) -> Dict[str, Any]:
        """
        List all documents in a collection

        Args:
            collection_name: Collection to list documents from

        Returns:
            List of documents with metadata
        """
        metadata_path = self.user_metadata_dir / f"{collection_name}.json"

        if not metadata_path.exists():
            return {
                "success": False,
                "error": f"Collection '{collection_name}' does not exist"
            }

        try:
            with open(metadata_path, 'r', encoding='utf-8') as f:
                metadata = json.load(f)

            documents = []
            for doc_id, doc_meta in metadata.get("documents", {}).items():
                chunk_count = len(self._get_document_chunks(doc_meta)) or len(doc_meta.get("chunk_indices", []))
                documents.append({
                    "id": doc_id,
                    "name": doc_meta["name"],
                    "path": doc_meta["path"],
                    "chunks": chunk_count,
                    "uploaded_at": doc_meta["uploaded_at"]
                })

            return {
                "success": True,
                "collection_name": collection_name,
                "documents": documents,
                "total_documents": len(documents),
                "total_chunks": metadata.get("chunk_count", 0)
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"Error listing documents: {str(e)}"
            }

    def delete_document(self, collection_name: str, document_id: str) -> Dict[str, Any]:
        """
        Delete a specific document from a collection
        Note: This rebuilds the FAISS index without the deleted document

        Args:
            collection_name: Collection containing the document
            document_id: ID of document to delete

        Returns:
            Result dictionary
        """
        metadata_path = self.user_metadata_dir / f"{collection_name}.json"

        if not metadata_path.exists():
            return {
                "success": False,
                "error": f"Collection '{collection_name}' does not exist"
            }

        try:
            # Load metadata
            with open(metadata_path, 'r', encoding='utf-8') as f:
                metadata = json.load(f)

            # Check if document exists
            if document_id not in metadata["documents"]:
                return {
                    "success": False,
                    "error": f"Document '{document_id}' not found in collection"
                }

            # Get document info before deletion
            deleted_doc = metadata["documents"][document_id]
            deleted_doc_name = deleted_doc["name"]
            deleted_doc_chunks = self._get_document_chunks(deleted_doc)
            deleted_chunks_count = len(deleted_doc_chunks)

            # Remove document from metadata
            del metadata["documents"][document_id]
            chunks_file = deleted_doc.get("chunks_file")
            if chunks_file:
                chunks_path = Path(chunks_file)
                if chunks_path.exists():
                    chunks_path.unlink()

            # Rebuild FAISS index without deleted document
            self._load_embedding_model()

            # Collect all remaining chunks and embeddings
            all_chunks = []
            for doc_id, doc_meta in metadata["documents"].items():
                all_chunks.extend(self._get_document_chunks(doc_meta))

            if len(all_chunks) > 0:
                # Re-generate embeddings for all remaining documents
                embeddings = self.embedding_model.encode(
                    all_chunks,
                    batch_size=config.RAG_EMBEDDING_BATCH_SIZE,
                    show_progress_bar=False,
                    normalize_embeddings=True
                )

                # Create new index
                import faiss
                if config.RAG_INDEX_TYPE == "Flat":
                    if config.RAG_SIMILARITY_METRIC == "cosine":
                        index = faiss.IndexFlatIP(self.embedding_dim)
                    else:
                        index = faiss.IndexFlatL2(self.embedding_dim)
                else:
                    index = faiss.IndexFlatL2(self.embedding_dim)

                # Add all embeddings
                index.add(np.array(embeddings).astype('float32'))

                # Update chunk indices for remaining documents
                current_idx = 0
                for doc_id, doc_meta in metadata["documents"].items():
                    chunk_count = len(self._get_document_chunks(doc_meta))
                    doc_meta["chunk_indices"] = list(range(current_idx, current_idx + chunk_count))
                    current_idx += chunk_count

                # Save updated index
                self._save_index(index, collection_name)
                metadata["chunk_count"] = index.ntotal
                _rebuild_chunk_lookup(metadata)
            else:
                # No documents left - create empty index
                import faiss
                index = faiss.IndexFlatL2(self.embedding_dim)
                self._save_index(index, collection_name)
                metadata["chunk_count"] = 0
                metadata["chunk_lookup"] = {}

            # Save updated metadata
            with open(metadata_path, 'w', encoding='utf-8') as f:
                json.dump(metadata, f, indent=2)

            return {
                "success": True,
                "collection_name": collection_name,
                "deleted_document": deleted_doc_name,
                "deleted_chunks": deleted_chunks_count,
                "remaining_documents": len(metadata["documents"]),
                "remaining_chunks": metadata["chunk_count"]
            }

        except Exception as e:
            return {
                "success": False,
                "error": f"Error deleting document: {str(e)}"
            }

    def _vram_mb(self) -> Optional[float]:
        """Return current GPU memory allocated in MB, or None if CUDA unavailable."""
        try:
            import torch
            if torch.cuda.is_available():
                return torch.cuda.memory_allocated() / (1024 ** 2)
        except ImportError:
            pass
        return None

    def _reset_peak_vram(self):
        try:
            import torch
            if torch.cuda.is_available():
                torch.cuda.reset_peak_memory_stats()
        except ImportError:
            pass

    def _peak_vram_mb(self) -> Optional[float]:
        try:
            import torch
            if torch.cuda.is_available():
                return torch.cuda.max_memory_allocated() / (1024 ** 2)
        except ImportError:
            pass
        return None

    def cleanup(self):
        """Clear instance references. Global model singletons are kept alive intentionally
        so the next request can reuse them without reloading from disk/GPU."""
        self.embedding_model = None

    def _get_document_chunks(self, doc_meta: Dict[str, Any]) -> List[str]:
        if "chunks" in doc_meta:
            return doc_meta["chunks"]

        chunks_file = doc_meta.get("chunks_file")
        if chunks_file:
            from tools.rag.memory_efficient_uploader import load_chunks_from_disk
            return load_chunks_from_disk(Path(chunks_file))

        return []

    def _chunk_text(self, text: str) -> List[str]:
        """
        Split text into chunks

        Args:
            text: Text to chunk

        Returns:
            List of chunks
        """
        chunk_size = config.RAG_CHUNK_SIZE
        overlap = config.RAG_CHUNK_OVERLAP

        chunks = []
        start = 0

        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]

            if chunk:
                chunks.append(chunk)

            start = end - overlap

        return chunks

    def _read_document(self, path: Path) -> str:
        """
        Read document content based on file type

        Args:
            path: Document path

        Returns:
            Document text
        """
        if path.suffix in ['.txt', '.md']:
            with open(path, 'r', encoding='utf-8') as f:
                return f.read()

        elif path.suffix == '.json':
            with open(path, 'r', encoding='utf-8') as f:
                return json.dumps(json.load(f), indent=2)

        elif path.suffix == '.csv':
            import pandas as pd
            df = pd.read_csv(path)
            return df.to_string()

        elif path.suffix in ['.xlsx', '.xls']:
            import pandas as pd
            excel_file = pd.ExcelFile(path)
            parts = []
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(path, sheet_name=sheet_name)
                parts.append(f"[Sheet: {sheet_name}]")
                parts.append(df.to_string())
            return '\n\n'.join(parts)

        elif path.suffix == '.pdf':
            # Use fast PDF reader
            try:
                from tools.rag.optimized_uploader import read_pdf_fast
                return read_pdf_fast(path)
            except Exception as e:
                print(f"[RAG] Fast PDF reader failed, falling back to PyPDFLoader: {e}")
                from langchain_community.document_loaders import PyPDFLoader
                loader = PyPDFLoader(str(path))
                pages = loader.load()
                return '\n'.join(page.page_content for page in pages)

        elif path.suffix == '.docx':
            from docx import Document
            doc = Document(path)
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    parts.append(' | '.join(cells))
            return '\n'.join(parts)

        else:
            with open(path, 'r', encoding='utf-8') as f:
                return f.read()

    def _load_or_create_index(self, collection_name: str, dim: int):
        """Load existing index or create new one"""
        import faiss

        index_path = self.user_index_dir / f"{collection_name}.index"

        if index_path.exists():
            return faiss.read_index(str(index_path))
        else:
            # Create new index based on config
            if config.RAG_INDEX_TYPE == "Flat":
                if config.RAG_SIMILARITY_METRIC == "cosine":
                    index = faiss.IndexFlatIP(dim)  # Inner product for cosine
                else:
                    index = faiss.IndexFlatL2(dim)
            else:
                # Default to Flat
                index = faiss.IndexFlatL2(dim)

            return index

    def _load_index(self, collection_name: str):
        """Load index — use process-level mtime cache to avoid disk reads per request."""
        import faiss

        index_path = self.user_index_dir / f"{collection_name}.index"
        if not index_path.exists():
            return None

        cache_key = f"{index_path}:{index_path.stat().st_mtime}"
        if cache_key in _GLOBAL_FAISS_CACHE:
            return _GLOBAL_FAISS_CACHE[cache_key]

        index = faiss.read_index(str(index_path))
        # Evict stale entries for this path before inserting the fresh one
        stale = [k for k in _GLOBAL_FAISS_CACHE if k.startswith(f"{index_path}:")]
        for k in stale:
            del _GLOBAL_FAISS_CACHE[k]
        _GLOBAL_FAISS_CACHE[cache_key] = index
        return index

    def _save_index(self, index, collection_name: str):
        """Save index"""
        import faiss

        # Ensure index directory exists
        self.user_index_dir.mkdir(parents=True, exist_ok=True)

        index_path = self.user_index_dir / f"{collection_name}.index"
        faiss.write_index(index, str(index_path))
    
    def _upload_pdf_optimized(
        self,
        collection_name: str,
        pdf_path: Path,
        document_name: Optional[str] = None,
        progress_callback: Optional[Any] = None
    ) -> Dict[str, Any]:
        """
        Upload PDF using optimized parallel processing
        
        Args:
            collection_name: Target collection
            pdf_path: Path to PDF file
            document_name: Optional override for document name
            progress_callback: Optional callback(message, progress_pct) for progress updates
            
        Returns:
            Upload result
        """
        from tools.rag.optimized_uploader import OptimizedRAGUploader
        
        self._load_embedding_model()
        
        # Create collection if it doesn't exist
        collection_dir = self.user_docs_dir / collection_name
        if not collection_dir.exists():
            collection_dir.mkdir(parents=True, exist_ok=True)
            print(f"[RAG] Created new collection directory: {collection_dir}")
        
        # Initialize optimized uploader
        uploader = OptimizedRAGUploader(
            embedding_model=self.embedding_model,
            embedding_dim=self.embedding_dim,
            max_workers=None,  # Use all CPU cores
            pages_per_batch=20  # Process 20 pages per worker
        )
        
        vram_before = self._vram_mb()
        self._reset_peak_vram()

        try:
            result = uploader.upload_pdf_optimized(
                pdf_path=pdf_path,
                collection_name=collection_name,
                user_docs_dir=self.user_docs_dir,
                user_index_dir=self.user_index_dir,
                user_metadata_dir=self.user_metadata_dir,
                document_name=document_name,
                progress_callback=progress_callback,
                show_progress_bar=True  # Enable text-based progress bar by default
            )
            if result.get("success") and vram_before is not None:
                vram_after = self._vram_mb()
                vram_peak = self._peak_vram_mb()
                result["vram_before_mb"] = round(vram_before, 1)
                result["vram_after_mb"] = round(vram_after, 1) if vram_after is not None else None
                result["vram_peak_mb"] = round(vram_peak, 1) if vram_peak is not None else None
                result["vram_delta_mb"] = round((vram_after - vram_before), 1) if vram_after is not None else None
            _release_cuda_cache("BaseRAGTool._upload_pdf_optimized")
            return result
            
        except Exception as e:
            print(f"[RAG] Optimized upload failed: {e}")
            print(f"[RAG] Falling back to standard upload method")
            
            # Fallback to standard method
            return self.upload_document(
                collection_name=collection_name,
                document_path=str(pdf_path),
                document_content=None,
                document_name=document_name,
                use_optimized=False  # Prevent infinite recursion
            )
